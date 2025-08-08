"""
Core components for BookWorm: DocumentProcessor, KnowledgeGraph, and MindmapGenerator
"""
import asyncio
import json
import logging
import os
import uuid
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple, Union
from dataclasses import dataclass, field
from datetime import datetime
import tempfile
import shutil

# Set up logger
logger = logging.getLogger(__name__)

# For PDF processing
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

# For document processing
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

# LightRAG components
try:
    from lightrag import LightRAG, QueryParam
    from lightrag.llm.openai import gpt_4o_mini_complete, openai_embed
    from lightrag.kg.shared_storage import initialize_pipeline_status
    from lightrag.utils import setup_logger as lightrag_setup_logger
except ImportError:
    LightRAG = None
    QueryParam = None
    gpt_4o_mini_complete = None
    openai_embed = None
    initialize_pipeline_status = None
    lightrag_setup_logger = None

from .utils import BookWormConfig, get_file_category, is_supported_file


@dataclass
class ProcessedDocument:
    """Data class for processed documents"""
    id: str = field(default_factory=lambda: str(uuid.uuid4()))
    original_path: str = ""
    processed_path: str = ""
    text_content: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)
    file_type: str = ""
    file_size: int = 0
    processed_at: datetime = field(default_factory=datetime.now)
    status: str = "pending"  # pending, processing, completed, failed
    error_message: Optional[str] = None


@dataclass 
class MindmapResult:
    """Data class for mindmap generation results"""
    document_id: str
    mermaid_syntax: str = ""
    html_content: str = ""
    markdown_outline: str = ""
    generated_at: datetime = field(default_factory=datetime.now)
    provider: str = ""
    token_usage: Dict[str, int] = field(default_factory=dict)


class DocumentProcessor:
    """Handles document ingestion and text extraction"""
    
    def __init__(self, config: BookWormConfig):
        self.config = config
        self.logger = logging.getLogger("bookworm.processor")
        
        # Ensure directories exist
        Path(self.config.document_dir).mkdir(parents=True, exist_ok=True)
        Path(self.config.processed_dir).mkdir(parents=True, exist_ok=True)
        
    async def process_document(self, file_path: Union[str, Path]) -> ProcessedDocument:
        """Process a single document and extract text"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if not is_supported_file(file_path):
            raise ValueError(f"Unsupported file type: {file_path.suffix}")
        
        # Check file size
        file_size = file_path.stat().st_size
        max_size = self.config.max_file_size_mb * 1024 * 1024
        if file_size > max_size:
            raise ValueError(f"File too large: {file_size} bytes (max: {max_size})")
        
        # Create document record
        doc = ProcessedDocument(
            original_path=str(file_path),
            file_type=get_file_category(file_path) or "unknown",
            file_size=file_size,
            status="processing"
        )
        
        try:
            # Extract text based on file type
            text_content = await self._extract_text(file_path)
            doc.text_content = text_content
            
            # Save processed text
            processed_filename = f"{doc.id}_{file_path.stem}.txt"
            processed_path = Path(self.config.processed_dir) / processed_filename
            
            with open(processed_path, 'w', encoding='utf-8') as f:
                f.write(text_content)
            
            doc.processed_path = str(processed_path)
            doc.status = "completed"
            
            # Extract metadata
            doc.metadata = await self._extract_metadata(file_path, text_content)
            
            self.logger.info(f"Successfully extracted text from document: {file_path} ({len(text_content)} chars)")
            return doc
            
        except Exception as e:
            doc.status = "failed"
            doc.error_message = str(e)
            self.logger.error(f"Failed to process document {file_path}: {e}")
            raise
    
    async def _extract_text(self, file_path: Path) -> str:
        """Extract text from various file formats"""
        file_category = get_file_category(file_path)
        
        if file_category == "pdf":
            return await self._extract_pdf_text(file_path)
        elif file_category == "text":
            return await self._extract_text_file(file_path)
        elif file_category == "document":
            return await self._extract_document_text(file_path)
        elif file_category in ["code", "data"]:
            return await self._extract_text_file(file_path)
        else:
            raise ValueError(f"Unsupported file category: {file_category}")
    
    async def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF files using mineru as primary processor (following user's methodology)"""
        # Try mineru first (user's preferred method from lightrag_ex.py)
        if self.config.pdf_processor == "mineru":
            try:
                return await self._extract_pdf_mineru(file_path)
            except Exception as e:
                logger.warning(f"MinerU extraction failed, falling back to PyMuPDF: {e}")
        
        # Fallback to other processors
        if self.config.pdf_processor == "pymupdf" and fitz:
            return await self._extract_pdf_pymupdf(file_path)
        elif self.config.pdf_processor == "pdfplumber" and pdfplumber:
            return await self._extract_pdf_pdfplumber(file_path)
        else:
            # Try available processors in order of preference
            try:
                return await self._extract_pdf_mineru(file_path)
            except Exception:
                if fitz:
                    return await self._extract_pdf_pymupdf(file_path)
                elif pdfplumber:
                    return await self._extract_pdf_pdfplumber(file_path)
                else:
                    raise ImportError("No PDF processing library available. Install mineru, pymupdf, or pdfplumber.")

    async def _extract_pdf_mineru(self, file_path: Path) -> str:
        """Extract PDF content using MinerU (following lightrag_ex.py methodology)"""
        try:
            # Import MinerU functions locally to avoid global import issues
            from mineru.cli.common import prepare_env, read_fn
            from mineru.data.data_reader_writer import FileBasedDataWriter
            from mineru.utils.enum_class import MakeMode
            from mineru.backend.pipeline.pipeline_analyze import doc_analyze as pipeline_doc_analyze
            from mineru.backend.pipeline.pipeline_middle_json_mkcontent import union_make as pipeline_union_make
            from mineru.backend.pipeline.model_json_to_middle_json import result_to_middle_json as pipeline_result_to_middle_json
            
            # Set environment variables to force CPU usage (like in user's script)
            import os
            import tempfile
            os.environ["CUDA_VISIBLE_DEVICES"] = ""
            
            # Create temporary directory for processing
            with tempfile.TemporaryDirectory() as temp_dir:
                # Prepare the output directory and file name
                pdf_file_name = file_path.stem
                local_image_dir, local_md_dir = prepare_env(temp_dir, pdf_file_name, "auto")
                
                # Read PDF file
                pdf_bytes = read_fn(str(file_path))
                if not pdf_bytes:
                    logger.error(f"Failed to read PDF file: {file_path}")
                    return ""
                
                # Process PDF using MinerU pipeline (matching user's methodology)
                pdf_bytes_list = [pdf_bytes]
                p_lang_list = ["en"]  # Default to English, can be made configurable
                
                # Analyze document with CPU-only mode
                logger.info("Processing PDF with MinerU CPU-only mode...")
                infer_results, all_image_lists, all_pdf_docs, lang_list, ocr_enabled_list = pipeline_doc_analyze(
                    pdf_bytes_list, p_lang_list, parse_method="auto", formula_enable=True, table_enable=True
                )
                
                if not infer_results:
                    logger.error(f"Failed to analyze PDF: {file_path}")
                    return ""
                
                # Process the first (and only) document
                model_list = infer_results[0]
                images_list = all_image_lists[0]
                pdf_doc = all_pdf_docs[0]
                _lang = lang_list[0]
                _ocr_enable = ocr_enabled_list[0]
                
                # Create data writers
                image_writer = FileBasedDataWriter(local_image_dir)
                
                # Convert to middle JSON format
                middle_json = pipeline_result_to_middle_json(
                    model_list, images_list, pdf_doc, image_writer, _lang, _ocr_enable, True
                )
                
                pdf_info = middle_json["pdf_info"]
                
                # Generate markdown content
                image_dir = str(os.path.basename(local_image_dir))
                markdown_content = pipeline_union_make(pdf_info, MakeMode.MM_MD, image_dir)
                
                # Ensure we return a string
                if isinstance(markdown_content, str):
                    return markdown_content
                elif isinstance(markdown_content, list):
                    return "\n".join(str(item) for item in markdown_content)
                else:
                    return str(markdown_content) if markdown_content else ""
                    
        except ImportError:
            logger.warning("MinerU not available, falling back to PyMuPDF")
            raise ImportError("MinerU not available")
        except Exception as e:
            logger.error(f"Error with MinerU PDF extraction: {e}")
            raise
    
    async def _extract_pdf_pymupdf(self, file_path: Path) -> str:
        """Extract text using PyMuPDF"""
        text_content = ""
        doc = fitz.open(str(file_path))
        
        for page_num in range(doc.page_count):
            page = doc[page_num]
            text_content += page.get_text() + "\n\n"
        
        doc.close()
        return text_content.strip()
    
    async def _extract_pdf_pdfplumber(self, file_path: Path) -> str:
        """Extract text using pdfplumber"""
        text_content = ""
        
        with pdfplumber.open(str(file_path)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content += page_text + "\n\n"
        
        return text_content.strip()
    
    async def _extract_text_file(self, file_path: Path) -> str:
        """Extract text from plain text files"""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        raise ValueError(f"Could not decode file {file_path} with any supported encoding")
    
    async def _extract_document_text(self, file_path: Path) -> str:
        """Extract text from office documents"""
        extension = file_path.suffix.lower()
        
        if extension in ['.docx', '.doc']:
            return await self._extract_docx_text(file_path)
        else:
            # For other document types, we'll need additional libraries
            # For now, return a placeholder
            return f"Document type {extension} not yet supported for text extraction"
    
    async def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX files"""
        if not DocxDocument:
            raise ImportError("python-docx not available. Install python-docx to process DOCX files.")
        
        doc = DocxDocument(str(file_path))
        text_content = ""
        
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        
        return text_content.strip()
    
    async def _extract_metadata(self, file_path: Path, text_content: str) -> Dict[str, Any]:
        """Extract metadata from file and content"""
        stat = file_path.stat()
        
        metadata = {
            "filename": file_path.name,
            "extension": file_path.suffix,
            "file_size": stat.st_size,
            "created_at": datetime.fromtimestamp(stat.st_ctime),
            "modified_at": datetime.fromtimestamp(stat.st_mtime),
            "word_count": len(text_content.split()),
            "char_count": len(text_content),
            "line_count": len(text_content.splitlines()),
        }
        
        return metadata
    
    async def process_directory(self, directory_path: Union[str, Path]) -> List[ProcessedDocument]:
        """Process all supported documents in a directory"""
        directory_path = Path(directory_path)
        
        if not directory_path.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        documents = []
        
        # Find all supported files
        supported_files = []
        for file_path in directory_path.rglob("*"):
            if file_path.is_file() and is_supported_file(file_path):
                supported_files.append(file_path)
        
        self.logger.info(f"Found {len(supported_files)} supported files in {directory_path}")
        
        # Process files with concurrency control
        semaphore = asyncio.Semaphore(self.config.max_concurrent_processes)
        
        async def process_single_file(file_path):
            async with semaphore:
                try:
                    return await self.process_document(file_path)
                except Exception as e:
                    self.logger.error(f"Failed to process {file_path}: {e}")
                    return None
        
        # Process all files
        tasks = [process_single_file(file_path) for file_path in supported_files]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Filter out failed results
        for result in results:
            if isinstance(result, ProcessedDocument):
                documents.append(result)
        
        self.logger.info(f"Successfully processed {len(documents)} documents")
        return documents


class KnowledgeGraph:
    """Handles LightRAG knowledge graph operations following user's methodology"""
    
    def __init__(self, config: BookWormConfig):
        self.config = config
        self.logger = logging.getLogger("bookworm.knowledge_graph")
        self.rag: Optional[LightRAG] = None
        self.is_initialized = False
        
        if not LightRAG:
            raise ImportError("LightRAG not available. Install lightrag-hku.")
    
    async def initialize(self) -> None:
        """Initialize LightRAG system following user's lightrag_manager.py methodology"""
        if self.is_initialized:
            self.logger.info("LightRAG Knowledge Graph already initialized")
            return
            
        self.logger.info("Initializing LightRAG Knowledge Graph...")
        
        # Ensure working directory exists
        Path(self.config.working_dir).mkdir(parents=True, exist_ok=True)
        
        # Try to import Ollama functions (following user's pattern), fallback to OpenAI
        ollama_available = False
        try:
            from lightrag.llm.ollama import ollama_model_complete, ollama_embed
            from lightrag.utils import EmbeddingFunc
            ollama_available = True
            self.logger.info("Using Ollama LLM functions")
        except ImportError:
            self.logger.warning("Ollama LLM functions not available, using OpenAI")
            try:
                from lightrag.llm.openai import gpt_4o_mini_complete, openai_embed
                ollama_available = False
            except ImportError:
                self.logger.error("Neither Ollama nor OpenAI LLM functions available")
                raise ImportError("No LLM functions available")
            
        # Debug configuration (like in user's lightrag_manager.py)
        self.logger.info("Configuration:")
        self.logger.info(f"  Working Directory: {self.config.working_dir}")
        self.logger.info(f"  LLM Model: {self.config.llm_model}")
        self.logger.info(f"  Embedding Model: {self.config.embedding_model}")
        
        # Initialize LightRAG with appropriate LLM backend
        if ollama_available:
            # Use Ollama (following user's pattern)
            self.rag = LightRAG(
                working_dir=self.config.working_dir,
                llm_model_func=ollama_model_complete,
                llm_model_name=self.config.llm_model,
                llm_model_kwargs={
                    "host": self.config.llm_host,
                    "options": {"num_ctx": 18192, "num_threads": 11},
                    "timeout": self.config.timeout,
                },
                embedding_func=EmbeddingFunc(
                    embedding_dim=self.config.embedding_dim,
                    max_token_size=self.config.max_embed_tokens,
                    func=lambda texts: self._robust_ollama_embed(
                        texts,
                        embed_model=self.config.embedding_model,
                        host=self.config.embedding_host
                    ),
                ),
                vector_storage="FaissVectorDBStorage",
                cosine_better_than_threshold=0.3,
            )
        else:
            # Fallback to OpenAI
            self.logger.info("Using OpenAI as LLM provider")
            self.rag = LightRAG(
                working_dir=self.config.working_dir,
                embedding_func=openai_embed,
                llm_model_func=gpt_4o_mini_complete,
            )
        
        # Initialize storage backends and pipeline
        await self.rag.initialize_storages()
        if initialize_pipeline_status:
            await initialize_pipeline_status()
        
        self.is_initialized = True
        self.logger.info("LightRAG Knowledge Graph initialized successfully")
    
    async def _robust_ollama_embed(self, texts, embed_model, host, max_retries=3, delay=1):
        """Robust wrapper for ollama_embed with retry logic (following user's pattern)"""
        from lightrag.llm.ollama import ollama_embed
        
        for attempt in range(max_retries):
            try:
                self.logger.debug(f"Embedding processing attempt {attempt + 1} with model {embed_model}")
                return await ollama_embed(
                    texts, 
                    embed_model=embed_model, 
                    host=host,
                    timeout=self.config.embedding_timeout,
                    options={"num_threads": 11}
                )
            except Exception as e:
                if attempt == max_retries - 1:
                    self.logger.error(f"Embedding failed after {max_retries} attempts: {e}")
                    raise
                
                self.logger.warning(f"Embedding attempt {attempt + 1} failed: {e}")
                self.logger.info(f"Retrying in {delay} seconds...")
                await asyncio.sleep(delay)
                delay *= 2  # Exponential backoff
    
    async def add_document(self, document: ProcessedDocument) -> None:
        """Add a processed document to the knowledge graph"""
        if not self.rag:
            raise RuntimeError("Knowledge graph not initialized. Call initialize() first.")
        
        if not document.text_content:
            raise ValueError("Document has no text content")
        
        try:
            # Insert document content into LightRAG
            self.logger.info(f"Adding document {document.id} to knowledge graph (this may take a while)...")
            await self.rag.ainsert(document.text_content)
            self.logger.info(f"✅ Successfully added document {document.id} to knowledge graph")
            
        except Exception as e:
            self.logger.error(f"❌ Failed to add document {document.id} to knowledge graph: {e}")
            raise
    
    async def query(self, query: str, mode: str = "hybrid", stream: bool = False, **kwargs) -> str:
        """Query the knowledge graph following user's lightrag_manager.py methodology"""
        if not self.is_initialized:
            await self.initialize()
        
        # Valid modes from user's lightrag_manager.py
        valid_modes = ["local", "global", "hybrid", "naive", "mix", "bypass"]
        if mode not in valid_modes:
            self.logger.warning(f"Invalid mode '{mode}'. Using 'hybrid' instead.")
            mode = "hybrid"
        
        try:
            self.logger.info(f"Querying with mode: {mode}")
            self.logger.debug(f"Question: {query}")
            
            if self.rag:
                resp = await self.rag.aquery(
                    query,
                    param=QueryParam(mode=mode, stream=stream, **kwargs),
                )
                return resp
            else:
                self.logger.error("RAG system not initialized")
                return ""
            
        except Exception as e:
            self.logger.error(f"Error during query: {e}")
            return ""
    
    async def batch_add_documents(self, documents: List[ProcessedDocument]) -> None:
        """Add multiple documents to the knowledge graph"""
        if not documents:
            return
        
        self.logger.info(f"Adding {len(documents)} documents to knowledge graph")
        
        for i, document in enumerate(documents, 1):
            try:
                await self.add_document(document)
                self.logger.info(f"Progress: {i}/{len(documents)} documents added")
                
            except Exception as e:
                self.logger.error(f"Failed to add document {document.id}: {e}")
                continue
        
        self.logger.info("Batch document addition completed")
    
    async def finalize(self) -> None:
        """Finalize and cleanup LightRAG resources"""
        if self.rag:
            await self.rag.finalize_storages()
            self.logger.info("LightRAG resources finalized")


class MindmapGenerator:
    """Generates mindmaps from documents using the integrated mindmap generator"""
    
    def __init__(self, config: BookWormConfig):
        self.config = config
        self.logger = logging.getLogger("bookworm.mindmap")
        
        # Ensure output directory exists
        Path(self.config.output_dir).mkdir(parents=True, exist_ok=True)
    
    async def generate_mindmap(self, document: ProcessedDocument) -> MindmapResult:
        """Generate a mindmap for a processed document"""
        if not document.text_content:
            raise ValueError("Document has no text content")
        
        self.logger.info(f"Generating mindmap for document {document.id}")
        
        try:
            # Here we would integrate the mindmap generator from the repo
            # For now, this is a placeholder implementation
            result = MindmapResult(
                document_id=document.id,
                provider=self.config.api_provider,
                mermaid_syntax=self._generate_placeholder_mindmap(document),
                html_content=self._generate_placeholder_html(document),
                markdown_outline=self._generate_placeholder_markdown(document)
            )
            
            # Save results to files
            await self._save_mindmap_files(result)
            
            self.logger.info(f"Mindmap generated successfully for document {document.id}")
            return result
            
        except Exception as e:
            self.logger.error(f"Failed to generate mindmap for document {document.id}: {e}")
            raise
    
    def _generate_placeholder_mindmap(self, document: ProcessedDocument) -> str:
        """Generate a placeholder Mermaid mindmap syntax"""
        # This is a simplified placeholder - in the full implementation,
        # this would use the actual mindmap generator logic
        return f"""mindmap
    ((📄 {document.metadata.get('filename', 'Document')}))
        ((📊 Overview))
            (Word Count: {document.metadata.get('word_count', 0)})
            (File Type: {document.file_type})
            (Size: {document.file_size} bytes)
        ((📝 Content Summary))
            (Generated from: {document.original_path})
            (Processed at: {document.processed_at})
"""
    
    def _generate_placeholder_html(self, document: ProcessedDocument) -> str:
        """Generate placeholder HTML for mindmap visualization"""
        mermaid_syntax = self._generate_placeholder_mindmap(document)
        
        return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>BookWorm Mindmap - {document.metadata.get('filename', 'Document')}</title>
    <script src="https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.min.js"></script>
</head>
<body>
    <div id="mermaid">
        <pre class="mermaid">
{mermaid_syntax}
        </pre>
    </div>
    <script>
        mermaid.initialize({{
            startOnLoad: true,
            theme: 'default',
            mindmap: {{ useMaxWidth: true }}
        }});
    </script>
</body>
</html>"""
    
    def _generate_placeholder_markdown(self, document: ProcessedDocument) -> str:
        """Generate placeholder Markdown outline"""
        return f"""# {document.metadata.get('filename', 'Document')}

## Overview
- **Word Count**: {document.metadata.get('word_count', 0)}
- **File Type**: {document.file_type}
- **Size**: {document.file_size} bytes

## Content Summary
- **Generated from**: {document.original_path}
- **Processed at**: {document.processed_at}

## Document Structure
This is a placeholder outline. In the full implementation, this would contain
the actual hierarchical structure extracted from the document content.
"""
    
    async def _save_mindmap_files(self, result: MindmapResult) -> None:
        """Save mindmap files to output directory"""
        base_path = Path(self.config.output_dir) / f"mindmap_{result.document_id}"
        
        # Save Mermaid syntax
        mermaid_path = base_path.with_suffix(".mmd")
        with open(mermaid_path, 'w', encoding='utf-8') as f:
            f.write(result.mermaid_syntax)
        
        # Save HTML
        html_path = base_path.with_suffix(".html")
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(result.html_content)
        
        # Save Markdown
        markdown_path = base_path.with_suffix(".md")
        with open(markdown_path, 'w', encoding='utf-8') as f:
            f.write(result.markdown_outline)
        
        self.logger.info(f"Mindmap files saved: {base_path}")
    
    async def batch_generate_mindmaps(self, documents: List[ProcessedDocument]) -> List[MindmapResult]:
        """Generate mindmaps for multiple documents"""
        if not documents:
            return []
        
        self.logger.info(f"Generating mindmaps for {len(documents)} documents")
        results = []
        
        for i, document in enumerate(documents, 1):
            try:
                result = await self.generate_mindmap(document)
                results.append(result)
                self.logger.info(f"Progress: {i}/{len(documents)} mindmaps generated")
                
            except Exception as e:
                self.logger.error(f"Failed to generate mindmap for document {document.id}: {e}")
                continue
        
        self.logger.info("Batch mindmap generation completed")
        return results
